import os
import uuid
from pathlib import Path
import fitz # PyMuPDF
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from app.models.document import Document
from app.schemas.upload_chat import DocumentUploadResponse
from app.core.config import settings
from google import genai
from google.genai import types

UPLOAD_DIR = Path(__file__).resolve().parents[2] / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

class DocumentService:
    def __init__(self):
        api_key = settings.GEMINI_API_KEY or "DUMMY_KEY_FOR_TESTING"
        self.client = genai.Client(api_key=api_key)

    def process_upload(self, user_uid: str, file: UploadFile, db: Session) -> DocumentUploadResponse:
        doc_id = str(uuid.uuid4())
        original_filename = Path(file.filename or "").name
        ext = os.path.splitext(original_filename)[1].lower()
        if ext not in [".pdf", ".docx"]:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")
            
        filepath = UPLOAD_DIR / f"{doc_id}{ext}"
        try:
            with filepath.open("wb") as buffer:
                bytes_written = 0
                while chunk := file.file.read(1024 * 1024):
                    bytes_written += len(chunk)
                    if bytes_written > MAX_UPLOAD_BYTES:
                        raise HTTPException(status_code=400, detail="File size exceeds the 10MB limit.")
                    buffer.write(chunk)
        except Exception:
            filepath.unlink(missing_ok=True)
            raise
            
        text = ""
        pages = 0
        
        try:
            if ext == ".pdf":
                doc = fitz.open(filepath)
                pages = len(doc)
                if pages > 300:
                    doc.close()
                    filepath.unlink(missing_ok=True)
                    raise HTTPException(status_code=400, detail="Document exceeds the maximum limit of 300 pages.")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            elif ext == ".docx":
                doc = docx.Document(filepath)
                pages = 1 # Approximation for DOCX
                for para in doc.paragraphs:
                    text += para.text + "\n"
        except HTTPException:
            raise
        except Exception as e:
            filepath.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=f"Failed to process document: {str(e)}")
            
        if not text.strip():
            filepath.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail="Document is empty or text could not be extracted.")
            
        # Generate Summary
        summary = self._generate_summary(text)
        
        db_doc = Document(
            id=doc_id,
            user_uid=user_uid,
            filename=original_filename,
            filepath=str(filepath),
            pages=pages,
            extracted_text=text,
            summary=summary
        )
        db.add(db_doc)
        db.commit()
        
        # Trigger AI Ingestion Pipeline
        from app.knowledge.ingestion import ingestion_pipeline
        metadata = {
            "document_id": doc_id,
            "document_type": "user_upload",
            "source_name": original_filename,
            "tenant_id": user_uid
        }
        try:
            # We run ingestion in the background or synchronously. For now, synchronously.
            ingestion_pipeline.process_document(text, metadata)
        except Exception as e:
            # We don't want ingestion failure to completely break upload
            print(f"Warning: Ingestion failed for {doc_id}: {e}")
        
        return DocumentUploadResponse(
            document_id=doc_id,
            filename=original_filename,
            pages=pages,
            summary=summary
        )

    def _generate_summary(self, text: str) -> str:
        prompt = "Summarize the following legal document in 2-3 concise sentences. Focus on the nature of the document and its primary purpose:\n\n" + text[:15000]
        
        import time
        import logging
        import concurrent.futures
        logger = logging.getLogger(__name__)
        
        max_retries = 4
        models_to_try = ['gemini-3.6-flash', 'gemini-3.6-pro']
        
        for attempt in range(max_retries):
            model_name = models_to_try[attempt % len(models_to_try)]
            try:
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(
                        self.client.models.generate_content,
                        model=model_name,
                        contents=prompt
                    )
                    response = future.result(timeout=15)
                return response.text.strip()
            except concurrent.futures.TimeoutError as e:
                logger.warning(f"Summary generation with {model_name} timed out after 15 seconds.")
                if attempt < max_retries - 1:
                    continue
            except Exception as e:
                logger.warning(f"Summary generation with {model_name} failed: {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
        
        return "This document was uploaded successfully, but an AI summary is temporarily unavailable. You can still ask questions about the document."

document_service = DocumentService()
